"""포렌식 멀티 에이전트 실행 스크립트

흐름:
    이미지 경로 입력
    → 사건 입력
    → 전략 수립 (HITL 승인)
    → 계획 수립 (HITL 승인)
    → Sub-Agent 실행
    → 결과 확인 (HITL: 보고서/추가분석/종료)
    → 보고서 저장

Usage:
    python scripts/run_agent.py
    python scripts/run_agent.py --verbose
"""

from __future__ import annotations

import argparse
import asyncio
import hashlib
import json
import logging
import sys
import time
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

import structlog

from agents.manager.graph import (
    ExecutionCallback,
    create_manager_state,
    run_execution,
    run_planning,
    run_report,
    run_strategy,
)
from config import LLMProvider, load_settings
from database.engine import get_engine, init_db
from disk_image_validator import validate_image_path
from llm_provider.anthropic import AnthropicProvider
from llm_provider.base import BaseLLMProvider
from llm_provider.openai import OpenAIProvider
from mcp_client.client import MCPClientManager


DIM = "\033[2m"
BOLD = "\033[1m"
RESET = "\033[0m"
GREEN = "\033[32m"
RED = "\033[31m"
YELLOW = "\033[33m"
CYAN = "\033[36m"
BLUE = "\033[34m"
CLEAR_LINE = "\033[2K\r"


def configure_logging(verbose: bool) -> None:
    """로그 레벨 설정"""
    level = logging.DEBUG if verbose else logging.WARNING
    structlog.configure(wrapper_class=structlog.make_filtering_bound_logger(level))
    logging.basicConfig(level=level, stream=sys.stderr)
    logging.getLogger("dissect").setLevel(level)
    logging.getLogger("mcp").setLevel(level)


async def async_input(prompt: str = "") -> str:
    """asyncio 이벤트 루프 내에서 UTF-8 stdin 읽기"""
    if prompt:
        sys.stdout.write(prompt)
        sys.stdout.flush()
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(
        None, lambda: sys.stdin.buffer.readline().decode("utf-8").strip()
    )


def _elapsed(start: float) -> str:
    """경과 시간 포맷팅"""
    s = int(time.time() - start)
    if s < 60:
        return f"{s}s"
    return f"{s // 60}m {s % 60}s"


def _truncate(text: str, max_len: int = 80) -> str:
    """텍스트를 max_len 이하로 잘라서 반환"""
    text = text.replace("\n", " ").strip()
    if len(text) <= max_len:
        return text
    return text[:max_len - 3] + "..."


def print_header(title: str) -> None:
    """구분선이 있는 헤더 출력"""
    width = 60
    print(f"\n{CYAN}{'━' * width}{RESET}")
    print(f"{CYAN}{BOLD}  {title}{RESET}")
    print(f"{CYAN}{'━' * width}{RESET}")


def print_section(title: str, content: str) -> None:
    """제목과 내용을 구분선으로 감싸서 출력"""
    print_header(title)
    print(content)
    print(f"{CYAN}{'━' * 60}{RESET}\n")


def print_phase(phase_num: int, title: str) -> None:
    """단계 시작 헤더 출력"""
    phases = ["", "전략 수립", "계획 수립", "Sub-Agent 실행", "결과 요약", "보고서 작성"]
    bar = ""
    for i in range(1, 6):
        if i < phase_num:
            bar += f" {GREEN}[{i}]{RESET}"
        elif i == phase_num:
            bar += f" {YELLOW}{BOLD}[{i}]{RESET}"
        else:
            bar += f" {DIM}[{i}]{RESET}"
    print(f"\n{bar}")
    print(f"  {BOLD}{title}{RESET}")


def print_step_progress(
    step_index: int,
    total_steps: int,
    step_name: str,
    agent_name: str,
    status: str = "running",
) -> None:
    """Sub-Agent 단계 진행 상황 출력

    Args:
        step_index: 현재 단계 (0-based)
        total_steps: 전체 단계 수
        step_name: 단계 이름
        agent_name: 담당 에이전트
        status: running, success, error, skip
    """
    icons = {
        "running": f"{YELLOW}  {RESET}",
        "success": f"{GREEN}  {RESET}",
        "error": f"{RED}  {RESET}",
        "skip": f"{DIM}  {RESET}",
    }
    icon = icons.get(status, "  ")
    progress = f"{DIM}[{step_index + 1}/{total_steps}]{RESET}"
    agent_badge = f"{BLUE}{agent_name}{RESET}"

    if status == "running":
        print(f"  {icon} {progress} {step_name} {DIM}({agent_badge}{DIM}){RESET}")
    else:
        print(f"  {icon} {progress} {step_name} {DIM}({agent_badge}{DIM}){RESET}")


def print_tool_call(tool_name: str, iteration: int) -> None:
    """도구 호출 출력"""
    short_name = tool_name.split("__")[-1] if "__" in tool_name else tool_name
    print(f"    {DIM}iter {iteration}: {short_name}{RESET}")


def print_tool_result(tool_name: str, is_error: bool, content_preview: str) -> None:
    """도구 실행 결과 출력"""
    short_name = tool_name.split("__")[-1] if "__" in tool_name else tool_name
    if is_error:
        print(f"    {RED}    {short_name}: {_truncate(content_preview, 60)}{RESET}")
    else:
        print(f"    {GREEN}    {short_name}{RESET} {DIM}{_truncate(content_preview, 50)}{RESET}")


def print_step_result(step_index: int, total_steps: int, step_name: str, agent_name: str, status: str, output: str, elapsed: str) -> None:
    """Sub-Agent 단계 완료 결과 출력

    Args:
        step_index: 현재 단계 (0-based)
        total_steps: 전체 단계 수
        step_name: 단계 이름
        agent_name: 담당 에이전트
        status: success 또는 error
        output: 출력 내용
        elapsed: 경과 시간
    """
    if status == "success":
        icon = f"{GREEN}  {RESET}"
    else:
        icon = f"{RED}  {RESET}"
    progress = f"{DIM}[{step_index + 1}/{total_steps}]{RESET}"
    agent_badge = f"{BLUE}{agent_name}{RESET}"
    print(f"  {icon} {progress} {step_name} {DIM}({agent_badge}, {elapsed}){RESET}")
    print(f"    {DIM}{_truncate(output, 70)}{RESET}")


def print_results_table(task_results: list) -> None:
    """실행 결과 요약 테이블 출력"""
    success = sum(1 for r in task_results if r.get("status") == "success")
    error = sum(1 for r in task_results if r.get("status") == "error")

    print(f"\n  {BOLD}실행 결과{RESET}: {GREEN}{success} 성공{RESET} / {RED}{error} 실패{RESET} / 총 {len(task_results)}단계")
    print()
    for r in task_results:
        if r.get("status") == "success":
            icon = f"{GREEN}+{RESET}"
        else:
            icon = f"{RED}x{RESET}"
        name = r.get("agent_name", "?")
        output = _truncate(r.get("output", ""), 60)
        print(f"  {icon} {BLUE}{name:>10}{RESET}  {output}")
    print()


def print_tools(tools: dict, verbose: bool = False) -> None:
    """발견된 MCP 도구 목록 출력"""
    print(f"  {BOLD}MCP 도구{RESET}: {len(tools)}개 발견")
    if verbose:
        for name, tool in tools.items():
            print(f"    {DIM}- {name}: {tool.description or '(설명 없음)'}{RESET}")
    if not tools:
        print(f"  {RED}사용 가능한 도구가 없습니다.{RESET}")


def create_llm_provider(settings, api_choice: str = "default") -> BaseLLMProvider:
    """설정에 따라 LLM 프로바이더 인스턴스 반환

    Args:
        settings: 애플리케이션 설정
        api_choice: "default" (기존 .env) 또는 "mindlogic" (MindLogic Gateway)
    """
    if api_choice == "mindlogic":
        from config import LLMConfig
        mindlogic_config = LLMConfig(
            provider=LLMProvider.OPENAI,
            model=settings.mindlogic_model,
            base_url=settings.mindlogic_base_url,
        )
        return OpenAIProvider(mindlogic_config, api_key=settings.mindlogic_api_key)

    if settings.llm.provider == LLMProvider.OPENAI:
        return OpenAIProvider(settings.llm, api_key=settings.llm_api_key)
    return AnthropicProvider(settings.llm, api_key=settings.llm_api_key)


class ConsoleExecutionCallback:
    """Sub-Agent 실행 진행 상황을 콘솔에 시각적으로 출력"""

    def __init__(self) -> None:
        self._step_starts: dict[int, float] = {}

    def on_step_start(self, step_index: int, total: int, step: dict, agent_name: str) -> None:
        """단계 시작 시 진행 바와 단계 정보 출력"""
        self._step_starts[step_index] = time.time()
        name = step.get("name", step.get("purpose", ""))

        bar = self._progress_bar(step_index, total)
        print(f"\n  {bar}")
        print_step_progress(step_index, total, name, agent_name, "running")

    def on_step_done(self, step_index: int, total: int, step: dict, agent_name: str, result: dict) -> None:
        """단계 완료 시 결과 요약 출력"""
        elapsed = _elapsed(self._step_starts.get(step_index, time.time()))
        name = step.get("name", step.get("purpose", ""))
        status = result.get("status", "error")
        output = result.get("output", "")

        print_step_result(step_index, total, name, agent_name, status, output, elapsed)

    def on_step_skip(self, step_index: int, total: int, step: dict) -> None:
        """수동 단계 건너뛸 때 출력"""
        name = step.get("name", step.get("purpose", ""))
        print_step_progress(step_index, total, name, "manual", "skip")

    def _progress_bar(self, current: int, total: int) -> str:
        """텍스트 진행 바 생성"""
        width = 20
        filled = int(width * current / total) if total > 0 else 0
        bar = f"{GREEN}{'█' * filled}{RESET}{DIM}{'░' * (width - filled)}{RESET}"
        pct = int(100 * current / total) if total > 0 else 0
        return f"{bar} {pct}% ({current}/{total})"


async def prompt_disk_image_path() -> tuple[str, str] | None:
    """유효한 디스크 이미지 경로를 입력받을 때까지 반복"""
    print("분석할 디스크 이미지 경로를 입력하세요. (지원 형식: E01, dd, raw)\n")
    while True:
        path_input = await async_input("Image path: ")
        if not path_input or path_input.lower() in ("quit", "exit", "q"):
            return None
        result = validate_image_path(path_input)
        if result.is_valid:
            assert result.format is not None
            return path_input.strip("'\""), result.format.value
        print(f"[Error] {result.error_message}\n")


def save_report(
    report: str,
    dfxml: str = "",
    output_dir: Path | None = None,
) -> tuple[Path, Path | None]:
    """보고서를 Word(.docx), DFXML을 별도 파일로 저장

    Args:
        report: 마크다운 형식 보고서 텍스트
        dfxml: DFXML XML 문자열 (비어있으면 저장하지 않음)
        output_dir: 출력 디렉터리 (None이면 docs/)

    Returns:
        (docx 파일 경로, dfxml 파일 경로 또는 None)
    """
    if output_dir is None:
        output_dir = Path(__file__).parent.parent / "docs"
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    doc = Document()
    for line in report.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.strip() == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph(line.strip())
            for run in p.runs:
                run.font.size = Pt(11)

    if dfxml:
        doc.add_page_break()
        doc.add_heading("부록: DFXML (Digital Forensics XML)", level=1)
        p = doc.add_paragraph(dfxml)
        for run in p.runs:
            run.font.size = Pt(9)

    docx_path = output_dir / f"report_{timestamp}.docx"
    doc.save(str(docx_path))

    dfxml_path = None
    if dfxml:
        dfxml_path = output_dir / f"report_{timestamp}.dfxml"
        dfxml_path.write_text(dfxml, encoding="utf-8")

    return docx_path, dfxml_path


async def run_strategy_hitl(state, llm) -> dict:
    """전략 수립 HITL 루프

    LLM이 전략을 생성하고 사용자가 승인/수정할 때까지 반복
    """
    print_phase(1, "전략 수립")
    current = state

    while True:
        start = time.time()
        print(f"  {DIM}LLM 호출 중...{RESET}", end="", flush=True)
        current = await run_strategy(current, llm)
        print(f"\r  {GREEN}전략 생성 완료{RESET} ({_elapsed(start)})")

        strategy = current.get("analysis_strategy", "")
        print_section("분석 전략 (조사 대상 아티팩트)", strategy)

        choice = (await async_input(f"  {BOLD}수락(y){RESET} / {BOLD}수정(n){RESET} > ")).lower()
        if choice == "y":
            print(f"  {GREEN}전략 확정{RESET}")
            return current

        feedback = ""
        while not feedback:
            feedback = await async_input("  수정 방향 > ")
            if not feedback:
                print(f"  {YELLOW}수정 방향을 반드시 입력해야 합니다.{RESET}")

        user_msg = current["messages"][0].get("content", "")
        current = create_manager_state(
            user_message=(
                f"{user_msg}\n\n[이전 전략]\n{strategy}\n\n"
                f"[수정 요청]: {feedback}\n위 수정 요청을 반영하여 새로운 분석 전략을 작성하세요."
            ),
            disk_image_path=current.get("disk_image_path"),
            disk_image_format=current.get("disk_image_format"),
            system_profile=current.get("system_profile"),
        )
        print(f"  {DIM}전략 재수립 중...{RESET}")


async def run_planning_hitl(state, llm, mcp) -> dict:
    """계획 수립 HITL 루프

    LLM이 계획을 생성하고 사용자가 승인/수정할 때까지 반복
    """
    print_phase(2, "계획 수립")
    current = state

    while True:
        start = time.time()
        print(f"  {DIM}LLM 호출 중...{RESET}", end="", flush=True)
        current = await run_planning(current, llm, mcp)
        print(f"\r  {GREEN}계획 생성 완료{RESET} ({_elapsed(start)})")

        plan = current.get("analysis_plan", "")
        steps = current.get("plan_steps", [])
        print_section("세부 실행 계획", plan)

        if not steps:
            print(f"  {YELLOW}실행 가능한 단계가 없습니다. 계획을 재수립합니다.{RESET}")
            feedback = await async_input("  계획에 포함할 내용 > ")
            user_msg = current["messages"][0].get("content", "")
            current = {
                **current,
                "messages": [{"role": "user", "content": f"{user_msg}\n\n[추가 요청]: {feedback}"}],
            }
            continue

        print(f"  {BOLD}파싱된 실행 단계: {len(steps)}개{RESET}")
        for j, s in enumerate(steps):
            name = s.get("name", s.get("purpose", ""))
            server = s.get("mcp_server", "none")
            print(f"    {DIM}{j+1}.{RESET} {name} {DIM}→ {BLUE}{server}{RESET}")

        choice = (await async_input(f"\n  {BOLD}수락(y){RESET} / {BOLD}수정(n){RESET} > ")).lower()
        if choice == "y":
            print(f"  {GREEN}계획 확정{RESET}")
            return current

        feedback = ""
        while not feedback:
            feedback = await async_input("  수정 방향 > ")
            if not feedback:
                print(f"  {YELLOW}수정 방향을 반드시 입력해야 합니다.{RESET}")

        user_msg = current["messages"][0].get("content", "")
        current = {
            **current,
            "messages": [
                {
                    "role": "user",
                    "content": (
                        f"{user_msg}\n\n[이전 계획]\n{plan}\n\n"
                        f"[수정 요청]: {feedback}\n위 수정 요청을 반영하여 새로운 계획을 작성하세요."
                    ),
                }
            ],
        }
        print(f"  {DIM}계획 재수립 중...{RESET}")


CACHE_DIR = Path(__file__).parent.parent / ".cache"


def _cache_key(image_path: str, user_input: str) -> str:
    """이미지 경로와 사건 입력으로 캐시 키 생성"""
    raw = f"{image_path}::{user_input}"
    return hashlib.sha256(raw.encode()).hexdigest()[:16]


def _save_cache(key: str, stage: str, data: dict) -> None:
    """전략 또는 계획 데이터를 JSON 파일로 저장

    Args:
        key: 캐시 키
        stage: "strategy" 또는 "planning"
        data: 캐시할 상태 필드
    """
    CACHE_DIR.mkdir(parents=True, exist_ok=True)
    path = CACHE_DIR / f"{key}_{stage}.json"
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2))


def _load_cache(key: str, stage: str) -> dict | None:
    """캐시 파일 로드, 없으면 None 반환

    Args:
        key: 캐시 키
        stage: "strategy" 또는 "planning"
    """
    path = CACHE_DIR / f"{key}_{stage}.json"
    if path.exists():
        return json.loads(path.read_text())
    return None


def parse_args() -> argparse.Namespace:
    """커맨드라인 인자 파싱"""
    parser = argparse.ArgumentParser(description="포렌식 멀티 에이전트 실행")
    parser.add_argument("--verbose", "-v", action="store_true", help="로그 출력")
    parser.add_argument(
        "--use-cache", action="store_true",
        help="동일 입력의 전략/계획 캐시 사용 (테스트용, 추후 삭제 예정)",
    )
    return parser.parse_args()


async def main() -> None:
    """멀티 에이전트 초기화 및 실행

    전체 흐름:
        1. 설정 로드 → DB/LLM/MCP 초기화
        2. 디스크 이미지 경로 입력
        3. 사건 개요 입력
        4. 전략 수립 HITL (승인/수정 반복)
        5. 계획 수립 HITL (승인/수정 반복)
        6. Sub-Agent 실행 (Dissect)
        7. 결과 요약 출력
        8. 게이트: r=보고서 생성 / a=추가 분석 / q=종료
    """
    args = parse_args()
    configure_logging(args.verbose)

    config_path = Path(__file__).parent.parent / "config" / "mcp_servers.json"
    settings = load_settings(config_path)

    print(f"\n  {BOLD}API 선택{RESET}:")
    print(f"    1. 기본 ({settings.llm.provider.value} / {settings.llm.model})")
    has_mindlogic = bool(settings.mindlogic_api_key)
    if has_mindlogic:
        print(f"    2. MindLogic Gateway ({settings.mindlogic_model})")
    else:
        print(f"    {DIM}2. MindLogic Gateway (MINDLOGIC_API_KEY 미설정){RESET}")

    api_choice = "default"
    choice = input(f"  선택 (1/2, 기본=1) > ").strip()
    if choice == "2" and has_mindlogic:
        api_choice = "mindlogic"
        print(f"  {GREEN}MindLogic Gateway 사용{RESET}: {settings.mindlogic_model}")
    else:
        print(f"  {GREEN}기본 API 사용{RESET}: {settings.llm.provider.value} / {settings.llm.model}")

    print(f"  {BOLD}MCP{RESET}:  {', '.join(settings.mcp.servers.keys()) or '(없음)'}")

    if api_choice == "default" and not settings.llm_api_key:
        print(f"  {RED}LLM_API_KEY가 설정되지 않았습니다.{RESET}")
        return
    if api_choice == "mindlogic" and not settings.mindlogic_api_key:
        print(f"  {RED}MINDLOGIC_API_KEY가 설정되지 않았습니다.{RESET}")
        return
    if not settings.database_url:
        print(f"  {RED}DATABASE_URL이 설정되지 않았습니다.{RESET}")
        return

    db_engine = get_engine(settings.database_url)
    await init_db(db_engine)
    print(f"  {BOLD}DB{RESET}:   초기화 완료")

    llm = create_llm_provider(settings, api_choice)

    async with MCPClientManager(settings.mcp) as mcp:
        tools = await mcp.list_tools()
        print_tools(tools, verbose=args.verbose)

        print_header("Forensic Multi-Agent System")

        image_result = await prompt_disk_image_path()
        if image_result is None:
            print(f"  {DIM}종료합니다.{RESET}")
            await db_engine.dispose()
            return
        image_path, image_format = image_result
        print(f"  {GREEN}이미지 등록{RESET}: {image_path} ({image_format.upper()})")

        system_profile = ""
        print(f"  {DIM}시스템 프로필 추출 중...{RESET}", end="", flush=True)
        try:
            start = time.time()
            profile_result = await mcp.call_tool(
                "dissect__extract_system_profile",
                {"image_path": image_path},
            )
            system_profile = mcp.get_tool_result_text(profile_result)
            print(f"\r  {GREEN}시스템 프로필 추출 완료{RESET} ({_elapsed(start)})")
        except Exception as exc:
            print(f"\r  {YELLOW}시스템 프로필 추출 실패: {exc}{RESET}")

        while True:
            print("\n사건 개요를 입력하세요. (quit으로 종료)")
            user_input = await async_input("사건 입력 > ")
            if not user_input or user_input.lower() in ("quit", "exit", "q"):
                print("종료합니다.")
                break

            try:
                state = create_manager_state(
                    user_message=user_input,
                    disk_image_path=image_path,
                    disk_image_format=image_format,
                    system_profile=system_profile,
                )

                cache_key = _cache_key(image_path, user_input) if args.use_cache else ""

                cached_strategy = _load_cache(cache_key, "strategy") if cache_key else None
                if cached_strategy:
                    state = {**state, **cached_strategy}
                    print(f"[Cache] 전략 캐시 적용 ({cache_key})")
                    print_section("분석 전략 (캐시)", state.get("analysis_strategy", ""))
                else:
                    state = await run_strategy_hitl(state, llm)
                    if cache_key:
                        _save_cache(cache_key, "strategy", {
                            "analysis_strategy": state.get("analysis_strategy", ""),
                            "messages": state.get("messages", []),
                            "phase": "planning",
                        })
                        print(f"[Cache] 전략 캐시 저장 ({cache_key})")

                accumulated_results = []

                while True:
                    cached_plan = _load_cache(cache_key, "planning") if cache_key else None
                    if cached_plan:
                        state = {**state, **cached_plan}
                        print(f"[Cache] 계획 캐시 적용 ({cache_key})")
                        print_section("세부 실행 계획 (캐시)", state.get("analysis_plan", ""))
                        print(f"  → 파싱된 실행 단계: {len(state.get('plan_steps', []))}개")
                        cache_key = ""
                    else:
                        state = await run_planning_hitl(state, llm, mcp)
                        if cache_key:
                            _save_cache(cache_key, "planning", {
                                "analysis_plan": state.get("analysis_plan", ""),
                                "plan_steps": state.get("plan_steps", []),
                                "messages": state.get("messages", []),
                                "phase": "execution",
                            })
                            print(f"[Cache] 계획 캐시 저장 ({cache_key})")
                            cache_key = ""

                    steps = state.get("plan_steps", [])
                    print_phase(3, f"Sub-Agent 실행 ({len(steps)}단계)")

                    exec_start = time.time()
                    cb = ConsoleExecutionCallback()
                    state = await run_execution(state, llm, mcp, callback=cb)

                    task_results = state.get("task_results", [])
                    accumulated_results.extend(task_results)

                    print(f"\n  {GREEN}실행 완료{RESET} ({_elapsed(exec_start)})")
                    print_results_table(task_results)

                    print_phase(4, "결과 요약")
                    start = time.time()
                    print(f"  {DIM}LLM 호출 중...{RESET}", end="", flush=True)
                    report_result = await run_report(state, llm)
                    print(f"\r  {GREEN}요약 완료{RESET} ({_elapsed(start)})")

                    summary = report_result.get("summary", "")
                    print_section("분석 결과 요약", summary)

                    gate = (
                        await async_input(
                            f"  {BOLD}보고서(r){RESET} / {BOLD}추가분석(a){RESET} / {BOLD}종료(q){RESET} > "
                        )
                    ).lower()

                    if gate == "r":
                        print_phase(5, "보고서 작성")
                        start = time.time()
                        print(f"  {DIM}LLM 호출 중...{RESET}", end="", flush=True)
                        state = {**state, "task_results": accumulated_results}
                        report_result = await run_report(state, llm)
                        print(f"\r  {GREEN}보고서 생성 완료{RESET} ({_elapsed(start)})")

                        report = report_result.get("report", "")
                        dfxml = report_result.get("dfxml", "")
                        print_section("포렌식 분석 보고서", report)

                        docx_path, dfxml_path = save_report(report, dfxml)
                        print(f"  {GREEN}보고서 저장{RESET}: {docx_path}")
                        if dfxml_path:
                            print(f"  {GREEN}DFXML 저장{RESET}: {dfxml_path}")
                        break
                    elif gate == "a":
                        additional = ""
                        while not additional:
                            additional = await async_input("  추가 조사 내용 > ")
                            if not additional:
                                print(f"  {YELLOW}조사할 내용을 입력해야 합니다.{RESET}")
                        state = {
                            **state,
                            "messages": [
                                {
                                    "role": "user",
                                    "content": f"{user_input}\n\n[추가 조사 요청]: {additional}",
                                }
                            ],
                            "plan_steps": [],
                            "task_results": [],
                        }
                        continue
                    else:
                        break

            except Exception as e:
                print(f"\n[Error] {e}\n")
                import traceback
                traceback.print_exc()

    await db_engine.dispose()


if __name__ == "__main__":
    asyncio.run(main())
